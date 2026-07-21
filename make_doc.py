# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

TARGET = 20

# (name, current_count)
data = [
    ("Smew", 1), ("Skunk", 2), ("Gopher", 3), ("Lord", 3), ("Newt", 3),
    ("Dua_Flipa", 4), ("Marley", 4), ("Bilby", 5), ("Millipede", 5),
    ("Ronnie", 5), ("Stella", 5), ("Felix", 6), ("Goose", 6), ("Shiraz", 6),
    ("Heligan", 7), ("Echo", 8), ("Jack", 8), ("Lars", 8), ("Polka", 8),
    ("Trixie", 8), ("Cricket", 9), ("Tiki", 9), ("Groot", 10), ("Pepe", 10),
    ("Sean", 11), ("Solyom", 11), ("Karen", 12), ("Luke", 12), ("Ray", 12),
    ("Toogie", 12), ("Petal", 13), ("Stark", 13), ("Lizzie", 14),
    ("Piggy", 14), ("Summer", 14), ("Puddle", 15), ("Therese", 15),
    ("Egg", 17), ("Lucy", 17), ("Bob", 18), ("Coco", 18), ("Beau", 19),
    ("Ernest", 19), ("Greyjoy", 19), ("Hanbury", 19),
]

# Band colours: read from Color_Bands1.jpg / Color_Bands2.jpg.
# Order in the table is B1 B2 B3 B4 B5 (outer -> middle -> outer). A diagonal glare/shadow
# across the middle of each photo hides many birds' "middle rings (B2-B4)"; the outer rings
# (B1/B5) and the cable tie are more reliable. Entries marked "(middle unclear, glare)" need
# checking against the original card or a re-shot photo. Birds absent from the card are marked
# with NOT_FOUND ("no matching colour found").
bands = {
    # —— Clearly readable ——
    "Ronnie":  "紫 绿 绿 紫 紫",
    "Marley":  "蓝 绿 绿（两端空）",
    "Polka":   "青 蓝 蓝 蓝 蓝",
    "Summer":  "粉 蓝 蓝 蓝 粉",
    "Jack":    "紫 粉 粉 粉 紫",
    "Lars":    "蓝 粉 粉 粉 蓝",
    "Echo":    "橙 红 红 红 橙",
    "Felix":   "黑 红 红 红 黑",
    "Heligan": "橙 粉 粉 粉 橙",
    "Lizzie":  "粉 粉 粉 粉（B1空）",
    "Groot":   "橙 橙 橙 橙",
    "Greyjoy": "灰 灰 灰 灰 灰",
    "Piggy":   "粉 紫 紫 紫 粉",
    "Karen":   "紫 紫 紫 紫 紫",
    # —— Outer rings / cable tie readable, middle unclear due to glare ——
    "Coco":    "红…红（中间约黄，欠清）",
    "Luke":    "外侧绿（中间反光不清）",
    "Puddle":  "蓝…黄（中间反光不清）",
    "Stark":   "外侧蓝/红（中间反光不清）",
    "Petal":   "外侧粉（中间反光不清）",
    "Shiraz":  "外侧黄（中间反光不清）",
    "Sean":    "外侧橙（中间反光不清）",
    "Cricket": "外侧灰（中间反光不清）",
    "Tiki":    "外侧黑（中间反光不清）",
    "Stella":  "外侧青 + 绑带绿（中间反光不清）",
    "Hanbury": "外侧黄 + 绑带红（中间反光不清）",
    "Therese": "外侧紫（中间反光不清）",
    "Bob":     "外侧黑 + 绑带黑（中间反光不清）",
    "Pepe":    "整体偏棕（疑反光，欠清）",
    "Egg":     "格内有备注文字+反光，欠清",
    "Ernest":  "色卡上近乎空白/反光，欠清",
}
NOT_FOUND = "没找到对应颜色"

doc = Document()

title = doc.add_heading("企鹅照片收集清单（少于 20 张的个体）", level=0)

p = doc.add_paragraph()
p.add_run("生成日期：2026-06-21    目标：每只个体至少补足到 ")
r = p.add_run(f"{TARGET}")
r.bold = True
p.add_run(" 张")

p = doc.add_paragraph()
p.add_run("共需补拍的个体数量：")
r = p.add_run(f"{len(data)} 只")
r.bold = True
total_short = sum(TARGET - c for _, c in data)
p.add_run("，合计至少需补拍 ")
r = p.add_run(f"{total_short} 张")
r.bold = True
p.add_run("。")

# ---- Shooting guidance ----
doc.add_heading("拍摄建议：只拍游客角度的全身正面照（不采集背面）", level=1)

p = doc.add_paragraph()
p.add_run("实验依据：").bold = True
p.add_run("全身照（exp1）测试准确率 0.950，明显高于肚子裁剪（exp2）的 0.866。"
          "因此身份信号不止在肚子——脸部花纹、胸前黑色横带、体型比例都是线索，"
          "裁太狠反而丢信息，且裁剪器误检会传播噪声。收集时以“全身正面”为准，不必强求肚子清晰完整。")

p = doc.add_paragraph()
p.add_run("只拍正面（重要）：").bold = True
p.add_run("企鹅背部是大面积均匀深色，个体之间几乎一样，背面几乎不含身份信息——从背面无法可靠识别。"
          "把正反一起训练会拉高类内方差、引入跨类混淆样本，反而拖累正面准确率。"
          "因此本次只采集正面照，不采集背面照；线上系统对“非正面”的照片应提示游客重拍，而不是硬给身份。")

bullets = [
    ("核心原则", "最终用户输入是“游客拍完的照片”，训练数据要匹配真实使用场景——"
                "以游客角度为主，不要只拍影棚级完美照。"),
    ("只拍正面", "只采集正面 / 3/4 侧前方的照片。不要采集纯背面照——背面对识别没有价值。"),
    ("框选要求", "整只企鹅入镜（头到脚），让 脸 + 胸前黑带 + 肚子斑点 尽量都在画面里。"
                "肚子被轻微遮挡不再是问题。"),
    ("避免采集", "纯背影、侧后方、身体被裁掉一截、整个正面被严重遮挡、太远太糊的照片。"),
    ("推荐拍法", "手机手持随手拍；眼平视或略微俯视（游客通常站着拍，机位略高于企鹅）；"
                "不同距离（近 / 中 / 远）、不同光照（逆光、阴影、正午硬光）；背景可杂乱、画面可有多只企鹅。"
                "一句话：拍“游客角度的全身正面 / 侧前方照”。"),
    ("对稀有个体（≤5 张）", "优先补满！只有 1–5 张的个体模型基本学不会。"
                "对这些个体先保证“正面全身清晰、能识别”，距离可放宽（近距离清晰照也行），把数量补上来最重要。"),
    ("对 6–19 张的个体", "以游客角度的正面照为主，重点增加角度 / 距离 / 光照的多样性，而不是重复拍同一个姿势。"),
]
for head, body in bullets:
    pp = doc.add_paragraph(style="List Bullet")
    rr = pp.add_run(head + "：")
    rr.bold = True
    pp.add_run(body)

# ---- Band-colour notes ----
doc.add_heading("脚环颜色（Color Bands）说明", level=1)
p = doc.add_paragraph()
p.add_run("数据来源：").bold = True
p.add_run("你提供的色卡照片 Color_Bands1.jpg / Color_Bands2.jpg。颜色顺序按表中 B1 B2 B3 B4 B5（外侧→中间→外侧），"
          "多数个体外侧两环(B1/B5)与绑带(cable tie)同色。")
p = doc.add_paragraph()
r = p.add_run("重要提醒：")
r.bold = True
p.add_run("两张色卡照片中部都有一条斜向反光/阴影，把一批个体的“中间环(B2–B4)”盖住了，这些标为“(中间反光不清)”。"
          "建议把色卡平铺、避免反光重拍一张，我就能把所有颜色读全读准。色卡上没有的个体写“"+NOT_FOUND+"”。")

# ---- Table ----
doc.add_heading("需补拍清单（按缺口从大到小排序，含脚环颜色）", level=1)

table = doc.add_table(rows=1, cols=5)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
for i, t in enumerate(["序号", "企鹅姓名", "目前数量", "至少需补拍", "脚环颜色 (B1–B5)"]):
    run = hdr[i].paragraphs[0].add_run(t)
    run.bold = True

# sort by shortfall desc (i.e. current asc), keep stable
rows = sorted(data, key=lambda x: x[1])
for idx, (name, cnt) in enumerate(rows, start=1):
    short = TARGET - cnt
    cells = table.add_row().cells
    cells[0].text = str(idx)
    cells[1].text = name
    cells[2].text = str(cnt)
    cells[3].text = f"+{short}"
    cells[4].text = bands.get(name, NOT_FOUND)

# total row
cells = table.add_row().cells
run = cells[1].paragraphs[0].add_run("合计")
run.bold = True
run = cells[2].paragraphs[0].add_run(str(sum(c for _, c in data)))
run.bold = True
run = cells[3].paragraphs[0].add_run(f"+{total_short}")
run.bold = True

doc.add_heading("优先级提示", level=1)
note = doc.add_paragraph()
note.add_run("建议明天的优先顺序：")
note.add_run("(1) 先补 ≤5 张的极稀有个体（Smew、Skunk、Gopher、Lord、Newt、Dua_Flipa、Marley、"
             "Bilby、Millipede、Ronnie、Stella）；(2) 再补 6–14 张的个体；"
             "(3) 最后处理 15–19 张、只差几张的个体。").bold = False

import os
out = r"C:\Users\14773\Desktop\企鹅照片收集清单.docx"
try:
    doc.save(out)
except PermissionError:
    out = r"C:\Users\14773\Desktop\企鹅照片收集清单_含脚环.docx"
    doc.save(out)
print("Saved:", out)
print("Penguins:", len(data), "Total shortfall:", total_short)
